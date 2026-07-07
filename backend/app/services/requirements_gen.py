"""Turn extracted document text into a list of functional requirements,
and render that list as a downloadable .docx."""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone

from app.services import llm

MAX_REQUIREMENTS = 40

# Output-format rules always appended so the reply parses into one requirement per line.
_FR_FORMAT = (
    "Produce a clean list of FUNCTIONAL requirements. Each must be atomic, testable, and "
    "phrased as 'The system shall ...'. Return ONLY the requirements, one per line, with no "
    "preamble, no numbering, and no markdown — just one requirement sentence per line."
)

_DEFAULT_PERSONA = "You are the Requirements Agent on the Werk delivery platform."
_SYSTEM = f"{_DEFAULT_PERSONA}\n{_FR_FORMAT}"


async def generate_requirements(
    source_text: str, instruction: str = "", system: str | None = None
) -> tuple[list[str], str]:
    """Return (requirements, source). `system` lets the caller supply the agent's
    own (possibly user-edited) instructions; the FR output format is always enforced."""
    text = (source_text or "").strip()
    if not text:
        return [], "heuristic"

    sys_prompt = f"{system.strip()}\n\n{_FR_FORMAT}" if system and system.strip() else _SYSTEM
    out = await _llm_requirements(text, instruction, sys_prompt)
    if out:
        return out[:MAX_REQUIREMENTS], "llm"
    return _heuristic_requirements(text)[:MAX_REQUIREMENTS], "heuristic"


async def _llm_requirements(text: str, instruction: str, system: str) -> list[str] | None:
    user = (
        (instruction.strip() + "\n\n" if instruction.strip() else "")
        + "Requirements-gathering summary:\n\n"
        + text[:12000]
    )
    out, provider = await llm.chat_complete(system, user, max_tokens=1500)
    if not out or provider == "none":
        return None
    return _split_lines(out)


def _split_lines(blob: str) -> list[str]:
    out: list[str] = []
    for raw in blob.splitlines():
        line = raw.strip()
        line = re.sub(r"^(FR[-\s]?\d+[:.)]?\s*|\d+[.)]\s*|[-*•]\s*)", "", line).strip()
        if len(line) >= 8:
            out.append(line)
    return out


# Lines that are clearly not requirements (slide markers, headers, page numbers).
_SLIDE = re.compile(r"^\[slide\b", re.IGNORECASE)
_SKIP = re.compile(
    r"^(agenda|overview|introduction|background|thank you|thanks|questions\??|next steps|"
    r"contents?|table of contents|summary|page \d+|\d+\s*$)",
    re.IGNORECASE,
)

# Subjects that are actors (people/roles) — phrased as "allow <actor> to ...".
_ACTORS = {
    "user", "users", "admin", "admins", "administrator", "administrators", "member", "members",
    "manager", "managers", "customer", "customers", "founder", "founders", "developer",
    "developers", "tester", "testers", "team", "teams", "operator", "operators", "guest",
    "guests", "owner", "owners", "client", "clients", "staff", "employee", "employees",
}


def _looks_like_title(line: str) -> bool:
    """Heading-ish: short, no terminal punctuation, no obvious requirement verb."""
    words = line.split()
    return len(words) <= 8 and not line.rstrip().endswith((".", "!", "?", ":"))


def _heuristic_requirements(text: str) -> list[str]:
    """Offline fallback: derive requirement-like statements from bullet/sentence lines."""
    candidates: list[str] = []
    prev_was_slide = False
    for raw in text.splitlines():
        stripped = raw.strip()
        if _SLIDE.match(stripped):
            prev_was_slide = True
            continue
        line = re.sub(r"^[-*•\d.)\s]+", "", stripped).strip()
        # the first line right after a slide marker is usually the slide title — skip if heading-ish
        if prev_was_slide and (_looks_like_title(line) or len(line) < 8):
            prev_was_slide = False
            continue
        prev_was_slide = False
        if not line or len(line) < 8 or _SKIP.match(line):
            continue
        for part in re.split(r"(?<=[.!?])\s+", line):
            p = part.strip().rstrip(".")
            if len(p) < 8 or _SKIP.match(p):
                continue
            candidates.append(_as_requirement(p))
    seen: set[str] = set()
    result: list[str] = []
    for c in candidates:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            result.append(c)
    return result


_MODAL = re.compile(
    r"^(?P<subj>[A-Za-z][A-Za-z ]{0,30}?)\s+(?:must|should|shall|can|will|may|need to|needs to)"
    r"\s+(?:be able to\s+)?(?P<rest>.+)$",
    re.IGNORECASE,
)


def _as_requirement(line: str) -> str:
    if line.lower().startswith(("the system shall", "the system must")):
        return line[0].upper() + line[1:]
    m = _MODAL.match(line)
    if m:
        subj = m.group("subj").strip().lower()
        rest = m.group("rest").strip()
        subj_word = subj.replace("the ", "").strip()
        if subj_word in _ACTORS:
            return f"The system shall allow {subj_word} to {rest}"
        # system/component subject → drop it, keep the action
        return f"The system shall {rest}"
    # imperative line, e.g. "Export reports to PDF"
    first = line[0].lower() + line[1:] if line else line
    return f"The system shall {first}"


def build_requirements_docx(title: str, requirements: list[str], source_name: str = "") -> bytes:
    """Render the requirements list as a .docx and return the bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    h = doc.add_heading(title or "Functional Requirements", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Generated by the Werk Requirements Agent · "
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        + (f" · Source: {source_name}" if source_name else "")
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_heading("Functional Requirements", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Requirement"
    for i, req in enumerate(requirements, start=1):
        cells = table.add_row().cells
        cells[0].text = f"FR-{i}"
        cells[1].text = req

    doc.add_paragraph()
    footer = doc.add_paragraph().add_run(
        f"{len(requirements)} requirement(s). Draft for review."
    )
    footer.italic = True
    footer.font.size = Pt(9)
    footer.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
